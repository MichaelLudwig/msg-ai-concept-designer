import streamlit as st
from docx import Document
from io import BytesIO


def replace_placeholder(doc, placeholder, replacement):
    # Platzhalter in Absätzen
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement)
    
    for section in doc.sections:
        # Platzhalter in Kopfzeilen
        header = section.header
        for paragraph in header.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)

def insert_chapter_at_placeholder(doc, placeholder, chapter):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Füge das Kapitel vor dem Platzhalter ein
            for element in chapter:
                new_paragraph = paragraph.insert_paragraph_before()
                new_paragraph.text = element['text']
                new_paragraph.style = element['style']
                if 'alignment' in element:
                    new_paragraph.alignment = element['alignment']
            # Entferne den Platzhalter-Text
            paragraph.text = paragraph.text.replace(placeholder, '')




def export_dokument_to_word (new_title,new_header,toc_list, content, glossar):
    
    # Bestehendes Dokument öffnen
    document = Document('IT-Konzept Template.docx')#---Plazhalter ersetzen------------------------------------------------------------------------------------
    replacements = {
        '{{Titel}}': new_header,
        '{{Titel2}}': "Pilotimplementierung KI Konzept Designer"
    }
    
    # Platzhalter ersetzen
    for placeholder, replacement in replacements.items():
        replace_placeholder(document, placeholder, replacement)

    # Kapitelstruktur einsetzen -----------------------------------------------------------------------------------
    # Definition der neuen Kapitel
    chapter_content = []
    for i, item in enumerate(toc_list):
        title_text = item["title"]
        help_text= item["help_text"]
        chapter_content.append({"text": title_text, "style": "Heading 1"})
        chapter_content.append({"text": "Hinweis", "style": "Hinweistext"})
        chapter_content.append({"text": help_text, "style": "Hinweistext"})
        
        lines = content[i].splitlines()
        for line in lines:
            if line.startswith("#### "):
                # Wenn eine neue Überschrift 2 gefunden wird
                heading = line.replace("#### ", "").strip()
                chapter_content.append({"text": heading, "style": "Heading 2"})
            else:
                if line.startswith("- **"):
                    # Wenn eine Aufzählung gefunden wird
                    point = line.replace("- **", "").replace("**", "").strip()
                    chapter_content.append({"text": point, "style": "List Bullet 3"})
                else:
                    line = line.replace("**", "").strip()
                    chapter_content.append({"text": line, "style": "Normal"})
        
        #chapter_content.append({"text": content[i], "style": "Normal"})

    #Glossar ergänzen
    chapter_content.append({"text": "Glossar", "style": "Heading 1"})
    if glossar:
        lines = glossar.splitlines()
        for line in lines:
            if line.startswith("### "):
                # Wenn eine neue Überschrift 2 gefunden wird
                heading = line.replace("### ", "").strip()
                chapter_content.append({"text": heading, "style": "Header"})
            else:
                if line.startswith("- **"):
                    # Wenn eine Aufzählung gefunden wird
                    point = line.replace("- **", "").replace("**", "").strip()
                    chapter_content.append({"text": point, "style": "List Bullet"})
                else:
                    line = line.replace("**", "").strip()
                    chapter_content.append({"text": line, "style": "Normal"})
        

    chapters = [
        {
            "placeholder": "{{Kapitelstruktur}}",
            "content": chapter_content
        }
    ]
    for chapter in chapters:
        insert_chapter_at_placeholder(document, chapter['placeholder'], chapter['content'])

    # Save the document to a BytesIO object
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    word_file = buffer
    # Provide the file as a download
    st.sidebar.download_button(
        label="Word Dokument herunterladen",
        data=word_file,
        file_name="Konzept.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    

